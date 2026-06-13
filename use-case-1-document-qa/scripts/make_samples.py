"""Generate a realistic demo corpus into sample-documents/.

Each artefact is built to mimic actual insurance paperwork in structure,
terminology and layout — so the RAG pipeline (extraction → chunking →
retrieval) is exercised the way it would be in production:

  - policy-gold-shield.pdf       12-page Gold Shield Health Policy
  - policy-silver-plus.pdf       12-page Silver Plus Health Policy
  - policy-platinum-elite.pdf    14-page Platinum Elite Health Policy
                                 (broader coverage → richer comparison)
  - claim-form-CF102-scan.png    CMS-1500-inspired claim form scan
  - medical-report-discharge.docx Hospital discharge summary

Each PDF includes: declarations cover page, table of contents, definitions,
schedule of benefits (table), exclusions, claim procedure, riders, contact
information, and signature block — same skeleton actual US/UK health policies
follow. Numbers/dates/IDs are fictitious but plausible.

Run:  python scripts/make_samples.py
"""
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from docx import Document
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "sample-documents"
OUT.mkdir(exist_ok=True)


# ── Shared style system ───────────────────────────────────────────────────────
def _styles():
    base = getSampleStyleSheet()
    cover_title = ParagraphStyle(
        "CoverTitle", parent=base["Title"], fontSize=24, leading=30,
        textColor=colors.HexColor("#0B2447"), spaceAfter=18, alignment=1,
    )
    cover_sub = ParagraphStyle(
        "CoverSub", parent=base["Normal"], fontSize=13, leading=16,
        textColor=colors.HexColor("#475569"), alignment=1, spaceAfter=10,
    )
    h1 = ParagraphStyle(
        "H1", parent=base["Heading1"], fontSize=16, leading=20,
        textColor=colors.HexColor("#0B2447"), spaceBefore=14, spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "H2", parent=base["Heading2"], fontSize=13, leading=16,
        textColor=colors.HexColor("#19376D"), spaceBefore=10, spaceAfter=6,
    )
    body = ParagraphStyle(
        "Body", parent=base["BodyText"], fontSize=10.5, leading=15,
        textColor=colors.HexColor("#0F172A"), spaceAfter=8, alignment=4,  # justify
    )
    small = ParagraphStyle(
        "Small", parent=base["BodyText"], fontSize=9, leading=12,
        textColor=colors.HexColor("#475569"),
    )
    return cover_title, cover_sub, h1, h2, body, small


def _table_style() -> TableStyle:
    return TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E0E7FF")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2447")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
    ])


# ── Realistic policy generator ────────────────────────────────────────────────
def _policy(
    filename: str, plan_name: str, plan_code: str, policy_number: str,
    effective_date: str, premium: str, deductible: str, oop_max: str,
    coinsurance: str, inpatient_limit: str, outpatient_limit: str,
    rx_limit: str, diagnostic_limit: str, dental: str, maternity: str,
    preex_clause: str, claim_window: str, customer_phone: str,
) -> None:
    cover_title, cover_sub, h1, h2, body, small = _styles()
    doc = SimpleDocTemplate(
        str(OUT / filename), pagesize=LETTER,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        title=f"{plan_name} Health Insurance Policy",
        author="Contoso Mutual Insurance Company",
    )
    story = []

    # ── Page 1 — Cover / declarations ────────────────────────────────────────
    story += [
        Spacer(1, 1.2 * inch),
        Paragraph("CONTOSO MUTUAL INSURANCE COMPANY", cover_title),
        Paragraph("Licensed under the Department of Insurance · NAIC #84317", cover_sub),
        Spacer(1, 0.6 * inch),
        Paragraph(f"<b>{plan_name}</b>", cover_title),
        Paragraph("Comprehensive Health Insurance Policy", cover_sub),
        Spacer(1, 0.5 * inch),
        Paragraph(f"<b>Plan Code:</b> {plan_code}", body),
        Paragraph(f"<b>Policy Number:</b> {policy_number}", body),
        Paragraph(f"<b>Effective Date:</b> {effective_date}", body),
        Paragraph("<b>Policy Term:</b> 12 months (auto-renewable)", body),
        Paragraph(f"<b>Annual Premium:</b> {premium}", body),
        Spacer(1, 0.6 * inch),
        Paragraph(
            "This Policy is a legal contract between the Insured and Contoso Mutual "
            "Insurance Company. Please read it carefully and retain it with your records. "
            "Coverage is provided in consideration of the application and the payment of "
            "premiums when due.", body,
        ),
        Spacer(1, 0.3 * inch),
        Paragraph("Underwritten and administered by Contoso Mutual Insurance Company", small),
        Paragraph("1200 Liberty Avenue, Suite 400, Wilmington, DE 19801 · contoso-mutual.example", small),
        PageBreak(),
    ]

    # ── Page 2 — Table of contents (text, not a Table — OCR-friendly) ───────
    toc_lines = [
        ("1.", "Definitions", "3"),
        ("2.", "Schedule of Benefits", "4"),
        ("3.", "Deductibles, Co-payments and Out-of-pocket Maximum", "5"),
        ("4.", "Eligibility and Enrollment", "6"),
        ("5.", "Pre-existing Conditions", "7"),
        ("6.", "Exclusions and Limitations", "8"),
        ("7.", "Claims Procedure", "9"),
        ("8.", "Coordination of Benefits", "10"),
        ("9.", "Optional Riders", "11"),
        ("10.", "Regulatory Notices and Contact Information", "12"),
    ]
    story += [Paragraph("Table of Contents", h1)]
    for n, title, page in toc_lines:
        story.append(Paragraph(f"<b>{n}</b> &nbsp; {title} &nbsp;&nbsp; .... &nbsp; page {page}", body))
    story.append(PageBreak())

    # ── Page 3 — Section 1: Definitions ──────────────────────────────────────
    story += [
        Paragraph("Section 1: Definitions", h1),
        Paragraph(
            "The following definitions apply throughout this Policy and any attached "
            "endorsements or riders:", body,
        ),
        Paragraph("1.1 Insured", h2),
        Paragraph(
            "Any individual named in the Schedule of Benefits as a primary insured, "
            "spouse, or dependent child, who is enrolled under this Policy and for whom "
            "the applicable premium has been paid in full.", body,
        ),
        Paragraph("1.2 Insurer", h2),
        Paragraph(
            "Contoso Mutual Insurance Company, the legal entity issuing this Policy. The "
            "Insurer shall not delegate its obligations under this contract without "
            "written notice to the Insured.", body,
        ),
        Paragraph("1.3 Deductible", h2),
        Paragraph(
            "The dollar amount the Insured must pay out of pocket each policy year for "
            "Eligible Expenses before this Policy begins to pay benefits, except for "
            "Preventive Care services which are not subject to the Deductible.", body,
        ),
        Paragraph("1.4 Outpatient Treatment", h2),
        Paragraph(
            "Medically necessary care, diagnostic services, surgery and follow-up "
            "rendered without requiring a continuous overnight admission to a hospital "
            "or other inpatient facility.", body,
        ),
        Paragraph("1.5 Eligible Expense", h2),
        Paragraph(
            "A charge for a service or supply that is (a) medically necessary, "
            "(b) provided by a licensed provider, (c) not excluded by Section 6, and "
            "(d) within the Usual, Customary and Reasonable charge in the geographic area.", body,
        ),
        PageBreak(),
    ]

    # ── Page 4 — Section 2: Schedule of Benefits ─────────────────────────────
    # NOTE: rendered as TEXT paragraphs (not a reportlab Table). Document
    # Intelligence's OCR is unreliable on thin-stroke reportlab table cells
    # — it hallucinates digits like "$20,000" → "da 30/geo". Text paragraphs
    # extract flawlessly. Same information, just a friendlier layout for OCR.
    schedule_items = [
        ("Inpatient Hospitalization (room, board, surgery)", inpatient_limit, coinsurance),
        ("Outpatient Treatment (clinic, specialist, day surgery)", outpatient_limit, coinsurance),
        ("Diagnostic Imaging (MRI, CT, ultrasound, X-ray)", diagnostic_limit, coinsurance),
        ("Laboratory and Pathology", "$10,000", coinsurance),
        ("Prescription Medication (formulary)", rx_limit, "30% co-pay"),
        ("Maternity and Newborn Care", maternity, coinsurance),
        ("Mental Health and Substance Use", "$25,000", coinsurance),
        ("Dental and Vision (basic)", dental, "50% co-pay"),
        ("Emergency Ambulance", "$3,000 per event", "0% co-payment (after Deductible)"),
        ("Preventive Care (annual physical, vaccines)",
         "100% covered", "0% co-payment (Deductible waived)"),
    ]
    story += [
        Paragraph("Section 2: Schedule of Benefits", h1),
        Paragraph(
            f"The annual benefit limits below apply per insured member per policy year "
            f"under the {plan_name} plan. Amounts in excess of the listed limits are the "
            f"responsibility of the Insured.", body,
        ),
        Paragraph("<b>2.1 Annual Coverage Limits</b>", h2),
    ]
    for i, (benefit, limit, share) in enumerate(schedule_items, start=1):
        story.append(Paragraph(
            f"<b>2.1.{i} {benefit}.</b> Annual Limit: <b>{limit}</b>. "
            f"Member share after Deductible: {share}.", body))
    story += [
        Spacer(1, 0.18 * inch),
        Paragraph(
            "All limits reset on each policy anniversary. Unused benefit balances do not "
            "carry forward into the next policy year.", small,
        ),
        PageBreak(),
    ]

    # ── Page 5 — Section 3: Deductible / co-pay / OOP max ────────────────────
    story += [
        Paragraph("Section 3: Deductibles, Co-payments and Out-of-pocket Maximum", h1),
        Paragraph("3.1 Annual Deductible", h2),
        Paragraph(
            f"The annual Deductible under the {plan_name} plan is "
            f"<b>{deductible}</b> per insured member. A family aggregate Deductible of "
            f"twice the individual amount applies when two or more family members are "
            f"covered. The Deductible accrues from Eligible Expenses incurred during the "
            f"policy year and resets on the anniversary date.", body,
        ),
        Paragraph("3.2 Co-insurance", h2),
        Paragraph(
            f"After the Deductible is satisfied, this Policy pays "
            f"<b>{coinsurance.replace(' co-payment','')}</b> of Eligible Expenses and the "
            f"Insured pays the remainder, subject to the Out-of-pocket Maximum.", body,
        ),
        Paragraph("3.3 Out-of-pocket Maximum", h2),
        Paragraph(
            f"The Out-of-pocket Maximum for the policy year is <b>{oop_max}</b> per "
            f"insured member. Once reached, this Policy pays 100% of Eligible Expenses "
            f"for the remainder of the policy year. Premiums, non-covered services and "
            f"amounts above Usual, Customary and Reasonable do not accrue toward this "
            f"maximum.", body,
        ),
        Paragraph("3.4 Network Differential", h2),
        Paragraph(
            "Use of in-network providers (listed at contoso-mutual.example/providers) is "
            "strongly recommended. Out-of-network services are subject to a 50% co-payment "
            "and balance billing by the provider.", body,
        ),
        PageBreak(),
    ]

    # ── Page 6 — Section 4: Eligibility ──────────────────────────────────────
    story += [
        Paragraph("Section 4: Eligibility and Enrollment", h1),
        Paragraph(
            "Coverage is available to individuals aged 18 to 64 who are legal residents "
            "of the United States. Dependents include a legal spouse or domestic partner "
            "and unmarried children under age 26.", body,
        ),
        Paragraph(
            "Open enrollment runs annually from November 1 through January 15. A Special "
            "Enrollment Period of 60 days applies following a qualifying life event such "
            "as marriage, birth or adoption of a child, involuntary loss of other "
            "coverage, or permanent relocation.", body,
        ),
        Paragraph(
            "Coverage becomes effective on the first day of the calendar month following "
            "receipt of the completed enrollment application and the first premium "
            "payment. The Insurer may decline coverage if material misrepresentations are "
            "discovered during underwriting.", body,
        ),
        Paragraph(
            "This Policy automatically renews each anniversary unless either party "
            "provides written notice of non-renewal at least 30 days before the renewal "
            "date.", body,
        ),
        PageBreak(),
    ]

    # ── Page 7 — Section 5: Pre-existing ─────────────────────────────────────
    story += [
        Paragraph("Section 5: Pre-existing Conditions", h1),
        Paragraph(
            "A <b>Pre-existing Condition</b> is any medical condition for which medical "
            "advice, diagnosis, care or treatment was recommended or received within the "
            "six-month period ending on the Insured's enrollment date.", body,
        ),
        Paragraph(preex_clause, body),
        Paragraph(
            "Conditions that are disclosed on the enrollment application and accepted by "
            "the Insurer in writing are deemed not to be Pre-existing Conditions for the "
            "purposes of this Section.", body,
        ),
        Paragraph(
            "Pregnancy in progress at the effective date of coverage is not considered a "
            "Pre-existing Condition. Newborn and adopted children added under Special "
            "Enrollment are likewise exempt.", body,
        ),
        PageBreak(),
    ]

    # ── Page 8 — Section 6: Exclusions ──────────────────────────────────────
    excl = [
        "Cosmetic surgery, including non-reconstructive plastic surgery and "
        "elective dermatological procedures.",
        "Experimental, investigational or unproven treatments not approved by the "
        "U.S. Food and Drug Administration for the condition being treated.",
        "Injuries sustained while participating in professional or semi-professional "
        "athletic competition.",
        "Self-inflicted injuries, except where resulting from a documented mental "
        "health condition covered under Section 2.",
        "Services rendered by an immediate family member of the Insured.",
        "Custodial care, long-term skilled nursing facility stays exceeding 100 days "
        "per policy year, and personal convenience items.",
        "War, declared or undeclared, and acts of terrorism in jurisdictions where "
        "the Insured is not lawfully present.",
        "Treatment of obesity unless meeting medical-necessity criteria for bariatric "
        "intervention as defined by the American Society for Metabolic and Bariatric "
        "Surgery.",
    ]
    story += [
        Paragraph("Section 6: Exclusions and Limitations", h1),
        Paragraph(
            "The following services, supplies and conditions are not covered under this "
            "Policy, except where specifically mandated by applicable law:", body,
        ),
    ]
    for i, item in enumerate(excl, start=1):
        story.append(Paragraph(f"<b>6.{i}</b> &nbsp; {item}", body))
    story.append(PageBreak())

    # ── Page 9 — Section 7: Claims procedure ─────────────────────────────────
    story += [
        Paragraph("Section 7: Claims Procedure", h1),
        Paragraph(
            f"<b>7.1</b> &nbsp; Claims for benefits under this Policy must be submitted "
            f"within <b>{claim_window}</b> of the date of service. Late submissions are "
            f"reviewed at the Insurer's discretion and may be denied except where the "
            f"Insured can demonstrate that timely submission was not reasonably possible.", body,
        ),
        Paragraph(
            "<b>7.2</b> &nbsp; All claims shall be submitted on Contoso claim form "
            "<b>CF-102</b> accompanied by (a) itemized invoices showing CPT or HCPCS "
            "codes, (b) a copy of the Insured's identification card, and (c) any "
            "supporting medical reports.", body,
        ),
        Paragraph(
            "<b>7.3</b> &nbsp; The Insurer will issue an Explanation of Benefits within "
            "10 business days of receipt of a complete claim and reimburse approved "
            "Eligible Expenses within 15 business days thereafter. Direct payment to "
            "participating providers is the default; reimbursement to the Insured "
            "applies where care was rendered out of network.", body,
        ),
        Paragraph(
            "<b>7.4</b> &nbsp; Pre-authorization is required for non-emergency inpatient "
            "admissions, surgeries scheduled more than 72 hours in advance, advanced "
            "imaging (MRI, CT, PET) and durable medical equipment over $1,000. Failure "
            "to obtain pre-authorization may result in a 25% benefit reduction.", body,
        ),
        Paragraph(
            "<b>7.5</b> &nbsp; Disputed claims may be appealed in writing within 180 days "
            "of the Explanation of Benefits. Final disputes are subject to binding "
            "arbitration administered by the American Arbitration Association.", body,
        ),
        PageBreak(),
    ]

    # ── Page 10 — Coordination of Benefits ───────────────────────────────────
    story += [
        Paragraph("Section 8: Coordination of Benefits", h1),
        Paragraph(
            "When the Insured is covered by more than one health benefit plan, this "
            "Policy coordinates benefits to ensure total reimbursement does not exceed "
            "100% of Eligible Expenses. Order of benefit determination follows the rules "
            "promulgated by the National Association of Insurance Commissioners.", body,
        ),
        Paragraph(
            "Where this Policy is the primary plan, benefits are paid as if no other "
            "coverage exists. Where it is the secondary plan, benefits are reduced by "
            "amounts payable under the primary plan. The Insured must report all other "
            "coverages at enrollment and within 30 days of any change.", body,
        ),
        Paragraph(
            "Subrogation: where a third party is legally liable for an injury or illness, "
            "the Insurer is subrogated to the Insured's right of recovery up to the "
            "amount of benefits paid under this Policy.", body,
        ),
        PageBreak(),
    ]

    # ── Page 11 — Riders ────────────────────────────────────────────────────
    story += [
        Paragraph("Section 9: Optional Riders", h1),
        Paragraph(
            "The following riders may be added at any policy anniversary, subject to "
            "additional premium and underwriting:", body,
        ),
        Paragraph("9.1 Critical Illness Rider", h2),
        Paragraph(
            "Lump-sum cash benefit of $50,000 payable on first diagnosis of cancer, "
            "myocardial infarction, stroke, end-stage renal failure or major organ "
            "transplant. 90-day waiting period applies from rider effective date.", body,
        ),
        Paragraph("9.2 Hospital Cash Rider", h2),
        Paragraph(
            "$200 per night of inpatient admission, payable from the second night "
            "onward, up to 90 nights per policy year. Paid in addition to Section 2 "
            "Inpatient Hospitalization benefits.", body,
        ),
        Paragraph("9.3 International Travel Rider", h2),
        Paragraph(
            "Extends emergency and urgent care coverage to incidents occurring outside "
            "the United States for trips up to 90 consecutive days. Excludes "
            "scheduled treatment sought abroad as 'medical tourism.'", body,
        ),
        PageBreak(),
    ]

    # ── Page 12 — Notices + contact ──────────────────────────────────────────
    story += [
        Paragraph("Section 10: Regulatory Notices and Contact Information", h1),
        Paragraph(
            "<b>10.1 State Mandates.</b> This Policy complies with the minimum benefit "
            "and consumer-protection standards of the state of issuance. Where state law "
            "requires broader benefits, those provisions apply notwithstanding any "
            "narrower language in this Policy.", body,
        ),
        Paragraph(
            "<b>10.2 Non-discrimination.</b> The Insurer does not deny, limit, condition "
            "or terminate coverage on the basis of race, color, national origin, age, "
            "disability, sex or gender identity.", body,
        ),
        Paragraph(
            "<b>10.3 Free-look Period.</b> The Insured may cancel this Policy within "
            "30 days of receipt for a full refund of premium, less any claims paid.", body,
        ),
        Paragraph("Contact Information", h2),
        Paragraph(f"<b>Member Services:</b> {customer_phone}", body),
        Paragraph("<b>Claims Inquiries:</b> 1-800-555-0142 · claims@contoso-mutual.example", body),
        Paragraph("<b>Pre-authorization:</b> 1-800-555-0185 · preauth@contoso-mutual.example", body),
        Paragraph("<b>Mailing Address:</b> Contoso Mutual Insurance, PO Box 12001, Wilmington DE 19801", body),
        Paragraph("<b>Online Portal:</b> members.contoso-mutual.example", body),
        Spacer(1, 0.45 * inch),
        Paragraph(
            "Signed for and on behalf of Contoso Mutual Insurance Company:", body,
        ),
        Spacer(1, 0.35 * inch),
        Paragraph(
            "_____________________________________ &nbsp;&nbsp;&nbsp; "
            "_____________________________________", body,
        ),
        Paragraph(
            "Dr. Eleanor Whitfield, Chief Underwriting Officer &nbsp;&nbsp;&nbsp;&nbsp;"
            "Marcus T. Vance, Corporate Secretary", small,
        ),
    ]

    doc.build(story)
    print(f"wrote {filename}  ({plan_name})")


# ── Realistic scanned-form image (CMS-1500-inspired) ──────────────────────────
def _claim_form_image() -> None:
    W, H = 1700, 2200
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    # Try a real font for realism; fall back gracefully
    def font(size):
        for candidate in ["arial.ttf", "Arial.ttf", "DejaVuSans.ttf"]:
            try:
                return ImageFont.truetype(candidate, size)
            except OSError:
                continue
        return ImageFont.load_default()

    title_f, label_f, value_f, small_f = font(38), font(22), font(28), font(18)

    # Header band
    draw.rectangle([(0, 0), (W, 130)], fill="#0B2447")
    draw.text((60, 35), "CONTOSO MUTUAL INSURANCE COMPANY", fill="white", font=title_f)
    draw.text((60, 90), "CLAIM FORM  ·  CF-102  ·  Health Insurance Reimbursement Request",
              fill="#A5D7E8", font=small_f)
    draw.text((W - 320, 50), "Form CF-102 (Rev. 2026-01)", fill="#CBD5E1", font=small_f)

    # Form sections with thin frames mimicking a printed form
    def field(x, y, w, h, label, value, *, value_font=value_f):
        draw.rectangle([(x, y), (x + w, y + h)], outline="#0B2447", width=2)
        draw.text((x + 10, y + 6), label, fill="#0B2447", font=label_f)
        draw.text((x + 14, y + 40), value, fill="#0F172A", font=value_font)

    def block(x, y, w, h, label):
        draw.rectangle([(x, y), (x + w, y + h)], outline="#0B2447", width=2)
        draw.text((x + 10, y + 6), label, fill="#0B2447", font=label_f)

    y = 170
    block(60, y, W - 120, 50, "SECTION A — CLAIMANT INFORMATION")
    y += 60
    field(60, y, 760, 86, "1. Member name (Last, First, MI)",  "Member, Jane Q.")
    field(840, y, 740, 86, "2. Member ID / Policy #",          "GS-2024-88811")
    y += 100
    field(60, y, 380, 86, "3. Date of Birth",                  "1986-04-22")
    field(460, y, 320, 86, "4. Sex",                            "F")
    field(800, y, 380, 86, "5. Relationship",                   "Self")
    field(1200, y, 380, 86, "6. Phone",                         "(415) 555-0177")
    y += 100
    field(60, y, 1520, 86, "7. Mailing address",                "228 Hawthorn Avenue, Apt 5B, San Francisco, CA 94117")
    y += 120

    block(60, y, W - 120, 50, "SECTION B — PROVIDER & SERVICE")
    y += 60
    field(60, y, 760, 86, "8. Provider / Facility name",        "Northwind Medical Center")
    field(840, y, 740, 86, "9. Provider Tax ID (TIN)",          "94-2847619")
    y += 100
    field(60, y, 760, 86, "10. Date of Service",                "14 March 2026")
    field(840, y, 740, 86, "11. Place of Service",              "Inpatient Hospital (POS 21)")
    y += 100
    field(60, y, 760, 86, "12. Diagnosis (ICD-10)",             "K35.80  Acute appendicitis, unspecified")
    field(840, y, 740, 86, "13. Procedure (CPT)",               "44970  Laparoscopic appendectomy")
    y += 120

    block(60, y, W - 120, 50, "SECTION C — CHARGES")
    y += 60
    field(60,   y, 380, 86, "14. Total billed",                 "$ 12,450.00")
    field(460,  y, 380, 86, "15. Deductible applied",           "$    500.00")
    field(860,  y, 380, 86, "16. Co-insurance (20%)",           "$  2,390.00")
    field(1260, y, 320, 86, "17. Claim amount",                 "$  9,560.00")
    y += 120

    block(60, y, W - 120, 50, "SECTION D — AUTHORIZATION")
    y += 60
    draw.text((75, y + 10),
              ("I certify that the information above is true and complete to the best of my "
               "knowledge and authorize Contoso Mutual"),
              fill="#0F172A", font=label_f)
    draw.text((75, y + 40),
              ("Insurance Company to release such information as is necessary to process "
               "this claim. I assign benefits to the named provider"),
              fill="#0F172A", font=label_f)
    draw.text((75, y + 70),
              "where indicated above.",
              fill="#0F172A", font=label_f)
    y += 130
    field(60,   y, 760, 100, "18. Member signature",            "Jane Q. Member", value_font=font(34))
    field(840,  y, 380, 100, "19. Signed on",                   "21 March 2026")
    field(1240, y, 340, 100, "20. Stamp",                       "RECEIVED 22-03-2026", value_font=small_f)

    y += 160
    draw.text((60, y),
              "For office use only — Claim ID: 2026-03-22-994127  ·  Adjudicator: A. Ramos  ·  Status: APPROVED",
              fill="#475569", font=small_f)

    # Slight skew + light noise so OCR has to do real work
    img = img.rotate(-0.7, fillcolor="white")
    # add scan-grain noise sparsely
    px = img.load()
    import random
    random.seed(42)
    for _ in range(2200):
        x = random.randint(0, W - 1); yy = random.randint(0, H - 1)
        g = random.randint(180, 245)
        px[x, yy] = (g, g, g)

    img.save(OUT / "claim-form-CF102-scan.png")
    print("wrote claim-form-CF102-scan.png  (CMS-1500-inspired)")


# ── Realistic discharge summary docx ──────────────────────────────────────────
def _medical_report() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run("NORTHWIND MEDICAL CENTER")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x0B, 0x24, 0x47)
    subtitle = doc.add_paragraph()
    r2 = subtitle.add_run("Department of General Surgery  ·  1450 Marina Boulevard, San Francisco, CA 94133  ·  Tel (415) 555-0199")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    doc.add_paragraph()

    h1 = doc.add_heading("Discharge Summary", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x0B, 0x24, 0x47)

    def kv(rows):
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Light Grid"
        for i, (k, v) in enumerate(rows):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v
        doc.add_paragraph()

    doc.add_heading("Patient Demographics", level=2)
    kv([
        ("Patient name", "Member, Jane Q."),
        ("Date of birth", "1986-04-22  (age 39)"),
        ("Sex", "Female"),
        ("Medical record number", "NMC-447120"),
        ("Insurance", "Contoso Mutual — Gold Shield Plan · Policy GS-2024-88811"),
        ("Admitting physician", "Aisha Karim, MD"),
        ("Attending surgeon", "Daniel Okafor, MD, FACS"),
        ("Admission date", "14 March 2026, 23:47 PST"),
        ("Discharge date", "16 March 2026, 11:20 PST"),
        ("Length of stay", "2 days"),
    ])

    doc.add_heading("Reason for Admission", level=2)
    doc.add_paragraph(
        "The patient was admitted via the Emergency Department with a 14-hour history "
        "of severe right lower quadrant abdominal pain, nausea and a single episode of "
        "non-bilious emesis. Clinical examination revealed localized tenderness with "
        "rebound at McBurney's point; Rovsing's sign positive. Vital signs on "
        "presentation: T 38.4 °C, HR 104, BP 122/78, SpO₂ 99% on room air."
    )

    doc.add_heading("Investigations", level=2)
    doc.add_paragraph(
        "Laboratory: WBC 14.8 ×10⁹/L with 82% neutrophils; CRP 96 mg/L; lactate 1.1 mmol/L; "
        "remainder of basic metabolic panel within normal limits.\n"
        "Imaging: Contrast-enhanced CT abdomen and pelvis demonstrated a dilated, "
        "non-compressible appendix measuring 11 mm with surrounding fat stranding; no "
        "evidence of perforation or abscess."
    )

    doc.add_heading("Diagnosis", level=2)
    doc.add_paragraph("Acute appendicitis without perforation (ICD-10 K35.80).")

    doc.add_heading("Clinical Course and Procedures", level=2)
    doc.add_paragraph(
        "The patient consented to and underwent laparoscopic appendectomy (CPT 44970) "
        "on 15 March 2026 at 02:10 PST. Intra-operative findings confirmed an "
        "inflamed, non-perforated appendix; histopathology subsequently demonstrated "
        "acute suppurative appendicitis. The procedure was uncomplicated; estimated "
        "blood loss was 20 mL. Postoperative recovery was uneventful: the patient "
        "tolerated oral fluids on the morning of postoperative day 1, ambulated "
        "independently, and pain was well controlled on a stepped-down regimen."
    )

    doc.add_heading("Medications on Discharge", level=2)
    doc.add_paragraph(
        "- Amoxicillin-clavulanate 875 mg orally twice daily for 7 days\n"
        "- Acetaminophen 1 g orally every 6 hours as needed for pain (max 4 g/day)\n"
        "- Ondansetron 4 mg orally every 8 hours as needed for nausea, for 48 hours"
    )

    doc.add_heading("Follow-up Instructions", level=2)
    doc.add_paragraph(
        "Outpatient surgical review with Dr. Okafor in 14 days. The patient is advised "
        "no heavy lifting (>5 kg) for 4 weeks and to avoid driving until cleared at "
        "follow-up. Patient should return to the Emergency Department promptly if she "
        "develops fever >38.5 °C, increasing abdominal pain, persistent vomiting, or "
        "wound discharge."
    )

    doc.add_heading("Insurance and Billing Notes", level=2)
    doc.add_paragraph(
        "Itemized hospital charges totalling $12,450.00 have been transmitted "
        "electronically to Contoso Mutual Insurance under claim CF-102 reference "
        "2026-03-22-994127. The patient has been counselled regarding her policy "
        "deductible ($500.00) and 20% co-insurance under the Gold Shield Schedule of "
        "Benefits."
    )

    sig = doc.add_paragraph()
    sig.add_run("\nElectronically signed: Daniel Okafor, MD, FACS  ·  License CA-A47391  ·  16 March 2026  17:42 PST").italic = True
    doc.save(OUT / "medical-report-discharge.docx")
    print("wrote medical-report-discharge.docx")


# ── Plan definitions (numbers tuned to produce real comparison questions) ─────
def main() -> None:
    _policy(
        filename="policy-gold-shield.pdf",
        plan_name="Gold Shield",
        plan_code="GS-HMO-2026",
        policy_number="GS-2024-88811",
        effective_date="01 January 2026",
        premium="$5,640 per year",
        deductible="$500 per insured member",
        oop_max="$6,000",
        coinsurance="80% paid by Policy / 20% member co-payment",
        inpatient_limit="$250,000",
        outpatient_limit="$20,000 per policy year",
        rx_limit="$5,000",
        diagnostic_limit="$8,000",
        dental="$1,500",
        maternity="$15,000",
        preex_clause=(
            "Pre-existing Conditions are covered after a continuous waiting period of "
            "<b>12 months</b> from the policy start date, provided the condition was "
            "fully and accurately disclosed at enrollment. During the waiting period, "
            "treatment for the Pre-existing Condition is at the Insured's expense; the "
            "Insurer will pay 50% of Eligible Expenses for emergency care arising from "
            "the Pre-existing Condition even during the waiting period."
        ),
        claim_window="90 days",
        customer_phone="1-800-555-0150 · members@contoso-mutual.example",
    )

    _policy(
        filename="policy-silver-plus.pdf",
        plan_name="Silver Plus",
        plan_code="SP-PPO-2026",
        policy_number="SP-2024-44502",
        effective_date="01 January 2026",
        premium="$3,420 per year",
        deductible="$1,500 per insured member",
        oop_max="$8,500",
        coinsurance="70% paid by Policy / 30% member co-payment",
        inpatient_limit="$150,000",
        outpatient_limit="$10,000 per policy year",
        rx_limit="$3,000",
        diagnostic_limit="$5,000",
        dental="$750",
        maternity="$10,000",
        preex_clause=(
            "Pre-existing Conditions are excluded from coverage for the first "
            "<b>24 months</b> of the policy and thereafter covered only with prior "
            "written approval from the Insurer. No waiver is available except where "
            "mandated by state law for guaranteed-issue groups."
        ),
        claim_window="60 days",
        customer_phone="1-800-555-0152 · members@contoso-mutual.example",
    )

    _policy(
        filename="policy-platinum-elite.pdf",
        plan_name="Platinum Elite",
        plan_code="PE-PPO-2026",
        policy_number="PE-2024-12077",
        effective_date="01 January 2026",
        premium="$9,180 per year",
        deductible="$250 per insured member",
        oop_max="$3,500",
        coinsurance="90% paid by Policy / 10% member co-payment",
        inpatient_limit="$1,000,000",
        outpatient_limit="$50,000 per policy year",
        rx_limit="$15,000",
        diagnostic_limit="$20,000",
        dental="$3,000 (includes vision)",
        maternity="$40,000",
        preex_clause=(
            "Pre-existing Conditions are covered from <b>day one</b> of this Policy "
            "with no waiting period, subject only to the Annual Deductible and "
            "Co-insurance. This day-one coverage applies to conditions fully disclosed "
            "at enrollment; intentional non-disclosure voids coverage for the "
            "concealed condition."
        ),
        claim_window="120 days",
        customer_phone="1-888-555-0190 · concierge@contoso-mutual.example",
    )

    _claim_form_image()
    _medical_report()
    print("\nDone — sample corpus written to:", OUT)


if __name__ == "__main__":
    main()
